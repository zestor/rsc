from __future__ import annotations

import csv
import io
import json
from dataclasses import dataclass
from pathlib import Path

from .observability import stable_hash, text_summary

TEXT_SUFFIXES = {
    ".csv",
    ".json",
    ".log",
    ".md",
    ".markdown",
    ".rst",
    ".text",
    ".txt",
    ".yaml",
    ".yml",
}


@dataclass(frozen=True)
class ConvertedDocument:
    filename: str
    media_type: str
    markdown: str
    source_format: str
    byte_count: int
    sha256: str

    def summary(self) -> dict:
        return {
            "filename": self.filename,
            "media_type": self.media_type,
            "source_format": self.source_format,
            "byte_count": self.byte_count,
            "sha256": self.sha256,
            "markdown": text_summary(self.markdown),
        }


@dataclass(frozen=True)
class DocumentChunk:
    chunk_id: str
    filename: str
    index: int
    total: int
    markdown: str
    word_count: int
    char_count: int

    def summary(self) -> dict:
        return {
            "chunk_id": self.chunk_id,
            "filename": self.filename,
            "index": self.index,
            "total": self.total,
            "word_count": self.word_count,
            "char_count": self.char_count,
            "markdown": text_summary(self.markdown),
        }


def convert_document_to_markdown(
    filename: str,
    content: bytes,
    *,
    media_type: str = "",
) -> ConvertedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf" or media_type == "application/pdf":
        markdown = _pdf_to_markdown(content)
        source_format = "pdf"
    elif suffix == ".docx" or media_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }:
        markdown = _docx_to_markdown(content)
        source_format = "docx"
    elif suffix == ".csv":
        markdown = _csv_to_markdown(content)
        source_format = "csv"
    elif suffix == ".json":
        markdown = _json_to_markdown(content)
        source_format = "json"
    elif suffix in TEXT_SUFFIXES or media_type.startswith("text/"):
        markdown = _text_to_markdown(content)
        source_format = suffix.removeprefix(".") or "text"
    else:
        raise ValueError(f"Unsupported attachment type for {filename}")
    return ConvertedDocument(
        filename=filename,
        media_type=media_type,
        markdown=markdown.strip(),
        source_format=source_format,
        byte_count=len(content),
        sha256=stable_hash(content.hex()),
    )


def documents_to_prompt_context(documents: list[ConvertedDocument]) -> str:
    if not documents:
        return ""
    sections = ["# Attached Document Context"]
    for index, document in enumerate(documents, start=1):
        sections.append(
            f"## Attachment {index}: {document.filename}\n"
            f"Format: {document.source_format}\n"
            f"SHA256: {document.sha256}\n\n"
            f"{document.markdown}"
        )
    return "\n\n".join(sections)


def chunk_converted_documents(
    documents: list[ConvertedDocument],
    *,
    target_chars: int = 12000,
    overlap_chars: int = 800,
) -> list[DocumentChunk]:
    chunks: list[DocumentChunk] = []
    for document in documents:
        pieces = _chunk_markdown(document.markdown, target_chars, overlap_chars)
        total = len(pieces)
        for index, piece in enumerate(pieces, start=1):
            chunks.append(
                DocumentChunk(
                    chunk_id=f"{Path(document.filename).stem or 'attachment'}-{index:04d}",
                    filename=document.filename,
                    index=index,
                    total=total,
                    markdown=piece,
                    word_count=len(piece.split()),
                    char_count=len(piece),
                )
            )
    return chunks


def _text_to_markdown(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def _chunk_markdown(markdown: str, target_chars: int, overlap_chars: int) -> list[str]:
    text = markdown.strip()
    if not text:
        return []
    if len(text) <= target_chars:
        return [text]
    blocks = []
    for block in _semantic_blocks(text):
        if len(block) > target_chars:
            blocks.extend(_split_long_text(block, target_chars))
        else:
            blocks.append(block)
    chunks: list[str] = []
    current: list[str] = []
    current_chars = 0
    for block in blocks:
        block_len = len(block)
        if current and current_chars + block_len > target_chars:
            chunk = "\n\n".join(current).strip()
            chunks.append(chunk)
            overlap = chunk[-overlap_chars:] if overlap_chars > 0 else ""
            current = [overlap, block] if overlap else [block]
            current_chars = len("\n\n".join(current))
        else:
            current.append(block)
            current_chars += block_len + 2
    if current:
        chunks.append("\n\n".join(current).strip())
    return chunks


def _semantic_blocks(text: str) -> list[str]:
    blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
    expanded: list[str] = []
    for block in blocks:
        if len(block) <= 14000:
            expanded.append(block)
            continue
        lines = block.splitlines()
        current: list[str] = []
        for line in lines:
            if len(line) > 8000:
                if current:
                    expanded.append("\n".join(current).strip())
                    current = []
                expanded.extend(_split_long_text(line, 8000))
                continue
            current.append(line)
            if sum(len(item) for item in current) > 8000:
                expanded.append("\n".join(current).strip())
                current = []
        if current:
            expanded.append("\n".join(current).strip())
    return expanded


def _split_long_text(text: str, size: int) -> list[str]:
    return [
        text[index : index + size].strip()
        for index in range(0, len(text), size)
        if text[index : index + size].strip()
    ]


def _csv_to_markdown(content: bytes) -> str:
    text = _text_to_markdown(content)
    rows = list(csv.reader(io.StringIO(text)))
    if not rows:
        return ""
    header = rows[0]
    body = rows[1:]
    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in body:
        padded = row + [""] * max(0, len(header) - len(row))
        lines.append("| " + " | ".join(padded[: len(header)]) + " |")
    return "\n".join(lines)


def _json_to_markdown(content: bytes) -> str:
    data = json.loads(_text_to_markdown(content))
    return f"```json\n{json.dumps(data, indent=2, sort_keys=True)}\n```"


def _pdf_to_markdown(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    sections = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            sections.append(f"## Page {index}\n{text.strip()}")
    return "\n\n".join(sections)


def _docx_to_markdown(content: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(content))
    lines = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
    return "\n\n".join(lines)
