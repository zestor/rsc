from __future__ import annotations

import logging
import queue
from io import BytesIO
from pathlib import Path

from fastapi.testclient import TestClient
from pypdf import PdfWriter

from rsc.config import RSCConfig
from rsc.contracts import EvalVerdict, LoopResult, LoopStatus, LoopTurnRecord
from rsc.document_converter import (
    ConvertedDocument,
    chunk_converted_documents,
    convert_document_to_markdown,
    documents_to_prompt_context,
)
import rsc.runtime as runtime
from rsc.web_api import (
    RunRequest,
    _QueueLogHandler,
    _ScopedSearchProvider,
    _activity_from_log,
    _planner_tasks_from_output,
    _bounded_join,
    _empty_attachment_task,
    _rubric_items_from_json,
    create_app,
)


def test_web_api_serves_react_static_and_health(tmp_path: Path):
    static_dir = tmp_path / "static"
    (static_dir / "assets").mkdir(parents=True)
    (static_dir / "index.html").write_text(
        "<div id='root'>RSC React UI</div><script src='/assets/app.js'></script>",
        encoding="utf-8",
    )
    (static_dir / "assets" / "app.js").write_text(
        "console.log('react bundle')", encoding="utf-8"
    )
    app = create_app(static_dir=static_dir)
    client = TestClient(app)

    health = client.get("/api/health")
    assert health.status_code == 200
    assert health.json() == {"ok": True, "ui_available": True}

    index = client.get("/")
    assert index.status_code == 200
    assert "RSC React UI" in index.text

    asset = client.get("/assets/app.js")
    assert asset.status_code == 200
    assert "react bundle" in asset.text


def test_web_api_run_endpoint_uses_injected_runner(tmp_path: Path):
    static_dir = tmp_path / "static"
    static_dir.mkdir()
    seen: list[RunRequest] = []

    def runner(request: RunRequest) -> dict:
        seen.append(request)
        return {
            "session_id": "session",
            "status": "passed",
            "final_score": 1.0,
            "turns_used": 1,
            "final_output": f"done: {request.task}",
            "total_tokens_input": 3,
            "total_tokens_output": 4,
            "memory_rules_added": [],
            "turns": [],
        }

    client = TestClient(create_app(runner=runner, static_dir=static_dir))
    response = client.post(
        "/api/runs",
        json={
            "task": "answer this",
            "rubric": [{"label": "complete", "description": "finish"}],
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["final_output"] == "done: answer this"
    assert seen[0].skill_name is None
    assert seen[0].rubric[0].label == "complete"


def test_web_api_config_endpoint_returns_safe_values(tmp_path: Path):
    config = RSCConfig(
        llm_provider="openrouter",
        openrouter_api_key="secret",
        loop_model="z-ai/glm-5.2",
        eval_model="z-ai/glm-5.2",
        state_dir=tmp_path / "state",
        log_dir=tmp_path / "logs",
        search_provider="none",
    )
    client = TestClient(
        create_app(static_dir=tmp_path / "static", config_factory=lambda: config)
    )
    response = client.get("/api/config")
    assert response.status_code == 200
    payload = response.json()
    assert payload["llm_provider"] == "openrouter"
    assert payload["loop_model"] == "z-ai/glm-5.2"
    assert payload["state_dir"] == str(tmp_path / "state")
    assert payload["skill_routing"].startswith("automatic")
    assert payload["skill_library_paths"] == []
    assert "secret" not in str(payload)


def test_web_api_config_model_options_adds_glm_fallback(tmp_path: Path):
    config = RSCConfig(
        llm_provider="openai",
        openai_api_key="secret",
        loop_model="gpt-5.5",
        eval_model="gpt-5.5",
        state_dir=tmp_path / "state",
        log_dir=tmp_path / "logs",
        search_provider="none",
    )
    client = TestClient(
        create_app(static_dir=tmp_path / "static", config_factory=lambda: config)
    )
    payload = client.get("/api/config").json()
    assert any(option["model"] == "gpt-5.5" for option in payload["models"])


def test_web_api_run_endpoint_converts_loop_result(tmp_path: Path):
    def runner(request: RunRequest) -> LoopResult:
        return LoopResult(
            session_id="session",
            task=request.task,
            final_output="finished",
            status=LoopStatus.PASSED,
            turns_used=1,
            final_score=1.0,
            turns=[
                LoopTurnRecord(
                    turn=1,
                    verdict=EvalVerdict(
                        passed=True, score=1.0, per_criterion={"complete": True}
                    ),
                )
            ],
            memory_rules_added=["rule"],
            total_tokens_input=1,
            total_tokens_output=2,
        )

    client = TestClient(create_app(runner=runner, static_dir=tmp_path / "static"))
    response = client.post("/api/runs", json={"task": "run"})
    assert response.status_code == 200
    assert response.json()["turns"][0]["verdict"]["passed"] is True


def test_web_api_run_endpoint_reports_runner_errors(tmp_path: Path):
    def runner(request: RunRequest) -> dict:
        del request
        raise RuntimeError("bad run")

    client = TestClient(create_app(runner=runner, static_dir=tmp_path / "static"))
    response = client.post("/api/runs", json={"task": "run"})
    assert response.status_code == 500
    assert response.json()["detail"] == "bad run"


def test_web_api_stream_endpoint_converts_files_and_streams_events(tmp_path: Path):
    seen: list[RunRequest] = []

    def runner(request: RunRequest) -> dict:
        seen.append(request)
        return {
            "session_id": "stream-session",
            "status": "passed",
            "final_score": 1.0,
            "turns_used": 1,
            "final_output": "streamed answer",
            "total_tokens_input": 5,
            "total_tokens_output": 6,
            "memory_rules_added": [],
            "turns": [],
        }

    client = TestClient(create_app(runner=runner, static_dir=tmp_path / "static"))
    response = client.post(
        "/api/runs/stream",
        data={
            "task": "Use the attachment",
            "mode": "research",
            "model": "z-ai/glm-5.2",
            "rubric_json": '[{"label":"complete","description":"finish"}]',
        },
        files={"files": ("note.md", b"# Attached\nImportant fact", "text/markdown")},
    )
    assert response.status_code == 200
    body = response.text
    assert "event: run.accepted" in body
    assert "event: attachments.convert.complete" in body
    assert "event: rsc.run.complete" in body
    assert "Attached Document Context" in seen[0].task
    assert "Important fact" in seen[0].task


def test_web_api_stream_endpoint_emits_errors(tmp_path: Path):
    def runner(request: RunRequest) -> dict:
        del request
        raise RuntimeError("stream failed")

    client = TestClient(create_app(runner=runner, static_dir=tmp_path / "static"))
    response = client.post(
        "/api/runs/stream",
        data={"task": "fail", "rubric_json": "[]"},
    )
    assert response.status_code == 200
    assert "event: rsc.run.error" in response.text
    assert "stream failed" in response.text


def test_web_api_stream_endpoint_uses_configured_orchestrator(
    tmp_path: Path, monkeypatch
):
    class FakeOrchestrator:
        def run(self, task, rubric, skill_name, session_id=None, **kwargs):
            assert (
                "Mode: custom" in task
                or "Mode: planning" in task
                or "Mode: research" in task
                or "Mode: synthesize" in task
            )
            assert rubric[0].label == "complete"
            assert skill_name == "default"
            return LoopResult(
                session_id=session_id or "configured",
                task=task,
                final_output="configured answer",
                status=LoopStatus.PASSED,
                turns_used=1,
                final_score=1.0,
            )

    seen_models: list[str] = []

    def fake_build_orchestrator(config):
        seen_models.append(config.loop_model)
        return FakeOrchestrator()

    monkeypatch.setattr("rsc.web_api.build_orchestrator", fake_build_orchestrator)
    client = TestClient(
        create_app(
            static_dir=tmp_path / "static",
            config_factory=lambda: RSCConfig(
                openrouter_api_key="key", search_provider="none"
            ),
        )
    )
    response = client.post(
        "/api/runs/stream",
        data={
            "task": "configured",
            "mode": "custom",
            "model": "custom-model",
            "rubric_json": '[{"label":"complete","description":"finish"}]',
        },
    )
    assert response.status_code == 200
    assert "event: rsc.run.complete" in response.text
    assert seen_models == ["custom-model"]


def test_web_api_stream_endpoint_decomposes_large_attachment(
    tmp_path: Path, monkeypatch
):
    calls: list[str] = []

    class FakeOrchestrator:
        def run(self, task, rubric, skill_name, session_id=None, **kwargs):
            del rubric, skill_name
            calls.append(task)
            return LoopResult(
                session_id=session_id or "final",
                task=task,
                final_output=f"result {len(calls)}",
                status=LoopStatus.PASSED,
                turns_used=1,
                final_score=1.0,
            )

    monkeypatch.setattr(
        "rsc.web_api.build_orchestrator", lambda config: FakeOrchestrator()
    )
    large = ("# Section\n\n" + "word " * 2600 + "\n\n") * 3
    client = TestClient(
        create_app(
            static_dir=tmp_path / "static",
            config_factory=lambda: RSCConfig(
                openrouter_api_key="key",
                search_provider="none",
                state_dir=tmp_path / "state",
            ),
        )
    )
    response = client.post(
        "/api/runs/stream",
        data={"task": "verify large document", "mode": "research", "rubric_json": "[]"},
        files={"files": ("large.md", large.encode("utf-8"), "text/markdown")},
    )
    assert response.status_code == 200
    assert "event: decomposition.plan" in response.text
    assert "event: chunk.verify.start" in response.text
    assert "event: chunk.verify.complete" in response.text
    assert len(calls) > 1
    assert "sequential document verification" in calls[0]
    assert "synthesize final answer" in calls[-1]
    run_dirs = list((tmp_path / "state" / "web_runs").glob("*"))
    assert len(run_dirs) == 1
    assert (run_dirs[0] / "manifest.json").exists()
    assert (run_dirs[0] / "section-summaries.md").exists()
    assert (run_dirs[0] / "final-synthesis-task.md").exists()
    chunk_dirs = list((run_dirs[0] / "chunks").glob("*"))
    assert chunk_dirs
    assert all((path / "source.md").exists() for path in chunk_dirs)
    assert all((path / "result.md").exists() for path in chunk_dirs)
    assert all((path / "summary.md").exists() for path in chunk_dirs)


def test_scoped_search_provider_uses_original_query_once():
    calls: list[tuple[str, int]] = []

    class Inner:
        name = "inner"

        def search(self, query: str, *, max_results: int = 5) -> str:
            calls.append((query, max_results))
            return f"results for {query}"

    provider = _ScopedSearchProvider(Inner(), "short user question")
    assert provider.name == "inner-scoped"
    assert (
        provider.search("specific chunk question", max_results=3)
        == "results for specific chunk question"
    )
    assert (
        provider.search("specific chunk question", max_results=9)
        == "results for specific chunk question"
    )
    assert (
        provider.search("x" * 600, max_results=5) == "results for short user question"
    )
    assert calls == [("specific chunk question", 3), ("short user question", 5)]


def test_document_converter_text_csv_json_and_unsupported():
    markdown = convert_document_to_markdown(
        "note.md", b"# Title", media_type="text/markdown"
    )
    assert markdown.markdown == "# Title"
    csv_doc = convert_document_to_markdown("data.csv", b"a,b\n1,2\n")
    assert "| a | b |" in csv_doc.markdown
    json_doc = convert_document_to_markdown("data.json", b'{"b": 2, "a": 1}')
    assert "```json" in json_doc.markdown
    context = documents_to_prompt_context([markdown, csv_doc])
    assert "Attachment 1: note.md" in context
    assert "Attachment 2: data.csv" in context
    try:
        convert_document_to_markdown("image.png", b"not supported")
    except ValueError as exc:
        assert "Unsupported attachment type" in str(exc)
    else:
        raise AssertionError("unsupported attachment should fail")


def test_document_converter_docx_pdf_empty_context_and_media_type():
    from docx import Document

    docx_buffer = BytesIO()
    doc = Document()
    doc.add_paragraph("Paragraph one")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    doc.save(docx_buffer)
    docx_doc = convert_document_to_markdown(
        "upload.bin",
        docx_buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "Paragraph one" in docx_doc.markdown
    assert "| A | B |" in docx_doc.markdown

    pdf_buffer = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.write(pdf_buffer)
    pdf_doc = convert_document_to_markdown(
        "blank.bin", pdf_buffer.getvalue(), media_type="application/pdf"
    )
    assert pdf_doc.source_format == "pdf"
    assert pdf_doc.markdown == ""
    assert convert_document_to_markdown("empty.csv", b"").markdown == ""
    assert documents_to_prompt_context([]) == ""


def test_document_chunking_splits_large_markdown_and_long_blocks():
    document = ConvertedDocument(
        filename="long.md",
        media_type="text/markdown",
        markdown=("alpha " * 2500) + "\n\n" + ("beta " * 2500),
        source_format="md",
        byte_count=1,
        sha256="hash",
    )
    chunks = chunk_converted_documents([document], target_chars=5000, overlap_chars=100)
    assert len(chunks) > 1
    assert max(chunk.char_count for chunk in chunks) <= 5200
    assert chunks[0].filename == "long.md"
    assert chunks[0].total == len(chunks)
    long_block = ConvertedDocument(
        filename="block.md",
        media_type="text/markdown",
        markdown="\n".join(["line " * 300 for _ in range(40)]),
        source_format="md",
        byte_count=1,
        sha256="hash",
    )
    assert len(chunk_converted_documents([long_block], target_chars=5000)) > 1


def test_friendly_activity_mapping_and_bad_rubric_json():
    search_start = _activity_from_log(
        {"event": "search.start", "query": {"preview": "latest topic"}}
    )
    assert search_start is not None
    assert search_start["data"]["title"] == "Searching the web"
    role_complete = _activity_from_log(
        {
            "event": "role.complete",
            "role": "planner",
            "output": {"preview": "Plan summary", "chars": 12},
        }
    )
    assert role_complete is not None
    assert role_complete["data"]["detail"] == "Plan summary"
    assert (
        _activity_from_log(
            {
                "event": "skill.selected",
                "skill_id": "research",
                "readiness": "ready",
                "score": 0.8,
            }
        )["data"]["title"]
        == "Selected skill context"
    )
    assert (
        _activity_from_log(
            {"event": "evaluator.complete", "passed": True, "score": 1.0}
        )["data"]["title"]
        == "Evaluation complete"
    )
    assert (
        _activity_from_log(
            {"event": "session.complete", "status": "passed", "final_score": 1.0}
        )["data"]["title"]
        == "Run complete"
    )
    assert _activity_from_log({"event": "unknown"}) is None
    assert "no extractable text" in _empty_attachment_task("question", "prepared")
    bounded = _bounded_join(["a" * 10, "b" * 10, "c" * 10], max_chars=25)
    assert "Omitted Section Summaries" in bounded
    try:
        _rubric_items_from_json("{}")
    except ValueError as exc:
        assert "rubric_json" in str(exc)
    else:
        raise AssertionError("non-list rubric_json should fail")


def test_queue_log_handler_emits_planner_tasks_and_activity():
    event_queue: queue.Queue[dict | None] = queue.Queue()
    handler = _QueueLogHandler(event_queue)
    record = logging.LogRecord(
        "rsc.test", logging.INFO, __file__, 1, "role.complete", (), None
    )
    record.event = "role.complete"
    record.session_id = "s"
    record.depth = 0
    record.fields = {
        "role": "planner",
        "turn": 2,
        "output": {
            "text": "## Steps\n1. Gather evidence.\n2. Verify chronology.",
            "preview": "## Steps",
            "chars": 44,
        },
    }
    handler.emit(record)
    first = event_queue.get_nowait()
    second = event_queue.get_nowait()
    third = event_queue.get_nowait()
    fourth = event_queue.get_nowait()
    assert first["event"] == "role.complete"
    assert first["data"]["role"] == "planner"
    assert "## Steps" in first["data"]["text"]
    assert second["event"] == "planner.tasks"
    assert second["data"]["turn"] == 2
    assert second["data"]["tasks"][0]["title"] == "Gather evidence."
    assert third["event"] == "activity.update"
    assert fourth["event"] == "log.role.complete"


def test_queue_log_handler_forwards_role_delta():
    event_queue: queue.Queue[dict | None] = queue.Queue()
    handler = _QueueLogHandler(event_queue)
    record = logging.LogRecord(
        "rsc.test", logging.INFO, __file__, 1, "role.delta", (), None
    )
    record.event = "role.delta"
    record.session_id = "s"
    record.depth = 0
    record.fields = {
        "role": "planner",
        "turn": 1,
        "sequence": 2,
        "delta": {"text": "streamed text", "preview": "streamed text", "chars": 13},
        "reasoning": {
            "text": "streamed reasoning",
            "preview": "streamed reasoning",
            "chars": 18,
        },
    }
    handler.emit(record)
    forwarded = event_queue.get_nowait()
    raw = event_queue.get_nowait()
    assert forwarded["event"] == "role.delta"
    assert forwarded["data"]["delta"]["preview"] == "streamed text"
    assert forwarded["data"]["reasoning"]["preview"] == "streamed reasoning"
    assert raw["event"] == "log.role.delta"


def test_planner_tasks_are_extracted_from_steps_section():
    output = """
## Approach
Short approach.

## Steps
1. Identify all factual claims and classify them by type.
2. Compare time-sensitive claims against known chronology.
3. Summarize unsupported or overstated conclusions.

## Risks
- Missing evidence.
"""
    tasks = _planner_tasks_from_output(output)
    assert [task["title"] for task in tasks] == [
        "Identify all factual claims and classify them by type.",
        "Compare time-sensitive claims against known chronology.",
        "Summarize unsupported or overstated conclusions.",
    ]


def test_web_api_missing_static_shows_build_message(tmp_path: Path):
    client = TestClient(create_app(static_dir=tmp_path / "missing"))
    response = client.get("/any/deep/link")
    assert response.status_code == 200
    assert "RSC UI is not built yet" in response.text


def test_runtime_build_search_provider_variants():
    firecrawl_config = RSCConfig(openai_api_key="key", search_provider="firecrawl")
    firecrawl = runtime.build_search_provider(firecrawl_config)
    assert firecrawl is not None
    assert firecrawl.name == "firecrawl"

    http_config = RSCConfig(
        openai_api_key="key",
        search_provider="http",
        search_endpoint="https://search.example",
    )
    http = runtime.build_search_provider(http_config)
    assert http is not None
    assert http.name == "http"

    none_config = RSCConfig(openai_api_key="key", search_provider="none")
    assert runtime.build_search_provider(none_config) is None


def test_runtime_build_client_variants(monkeypatch):
    class FakeOpenRouter:
        def __init__(self, **kwargs):
            self.kwargs = kwargs

    class FakeOpenAI:
        def __init__(self, **kwargs):
            self.kwargs = kwargs

    class FakeResponses:
        def __init__(self, client, **kwargs):
            self.client = client
            self.kwargs = kwargs

    monkeypatch.setattr(runtime, "OpenRouterClientAdapter", FakeOpenRouter)
    monkeypatch.setattr(runtime, "OpenAI", FakeOpenAI)
    monkeypatch.setattr(runtime, "OpenAIResponsesClientAdapter", FakeResponses)

    router_client = runtime.build_client(
        RSCConfig(
            llm_provider="openrouter",
            openrouter_api_key="router-key",
            loop_model="z-ai/glm-5.2",
        )
    )
    assert router_client.kwargs["api_key"] == "router-key"
    assert router_client.kwargs["model"] == "z-ai/glm-5.2"

    responses_client = runtime.build_client(
        RSCConfig(
            llm_provider="openai",
            openai_api_key="openai-key",
            openai_use_responses_api=True,
        )
    )
    assert responses_client.client.kwargs["api_key"] == "openai-key"
    assert responses_client.kwargs["text_verbosity"] == "medium"

    openai_client = runtime.build_client(
        RSCConfig(
            llm_provider="openai",
            openai_api_key="openai-key",
            openai_use_responses_api=False,
        )
    )
    assert openai_client.kwargs["api_key"] == "openai-key"


def test_runtime_build_orchestrator_wires_components(tmp_path: Path, monkeypatch):
    class FakeClient:
        model = "z-ai/glm-5.2"

        @property
        def chat(self):
            return self

        @property
        def completions(self):
            return self

        @property
        def embeddings(self):
            return self

    monkeypatch.setattr(runtime, "build_client", lambda config: FakeClient())
    monkeypatch.setattr(
        runtime, "configure_daily_file_logging", lambda log_dir, level: Path(log_dir)
    )
    config = RSCConfig(
        openrouter_api_key="key",
        search_provider="none",
        state_dir=tmp_path / "state",
        log_dir=tmp_path / "logs",
    )
    (config.state_dir / "skills").mkdir(parents=True)
    (config.state_dir / "memory_ledger.json").write_text(
        '{"schema_version":"1.0","entries":[]}'
    )
    orchestrator = runtime.build_orchestrator(config)
    assert orchestrator.model == "z-ai/glm-5.2"
    assert orchestrator.search_provider is None
